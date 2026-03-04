"""文档解析器模块

支持解析 doc, docx, pdf 格式的文档
"""

import sys
import os
import argparse
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
import json

# 模块级别编码配置 - 解决 Windows 控制台中文乱码问题
# 必须在任何其他代码之前执行
if sys.platform == 'win32':
    # 设置标准输出/错误编码为 UTF-8
    if sys.stdout.encoding != 'utf-8':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except Exception:
            # 如果无法重新配置，则设置环境变量
            os.environ['PYTHONIOENCODING'] = 'utf-8'
    if sys.stderr.encoding != 'utf-8':
        try:
            sys.stderr.reconfigure(encoding='utf-8')
        except Exception:
            pass

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class documentParser:
    """文档解析器，支持 doc, docx, pdf 格式"""

    def __init__(self, verbose: bool = False):
        self.supportedFormats = ['.docx', '.doc', '.pdf']
        self.verbose = verbose
        if verbose:
            logger.setLevel(logging.DEBUG)

    def parse(self, filePath: str) -> Dict[str, Any]:
        """
        解析文档文件

        Args:
            filePath: 文档文件路径

        Returns:
            包含解析结果的字典

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(filePath)

        if not path.exists():
            logger.error(f"文件不存在: {filePath}")
            raise FileNotFoundError(f"文件不存在: {filePath}")

        ext = path.suffix.lower()
        if ext not in self.supportedFormats:
            logger.error(f"不支持的文件格式: {ext}")
            raise ValueError(f"不支持的文件格式: {ext}. 支持的格式: {', '.join(self.supportedFormats)}")

        logger.info(f"开始解析文件: {path.name}, 格式: {ext}")

        # 根据文件格式选择解析方法
        if ext == '.docx':
            return self._parseDocx(path)
        elif ext == '.doc':
            return self._parseDoc(path)
        elif ext == '.pdf':
            return self._parsePdf(path)

    def _parseDocx(self, path: Path) -> Dict[str, Any]:
        """解析 docx 文件"""
        try:
            from docx import Document
        except ImportError:
            logger.error("缺少 python-docx 库")
            return {
                "success": False,
                "error": "请安装 python-docx 库: pip install python-docx",
                "format": "docx"
            }

        doc = Document(str(path))
        logger.debug(f"docx 文件段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")

        # 提取段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # 提取表格（增强版）
        tables = []
        for idx, table in enumerate(doc.tables):
            tableData = []
            for row in table.rows:
                rowData = [cell.text.strip() for cell in row.cells]
                tableData.append(rowData)

            if tableData:
                # 增强表格信息
                enhancedTable = self._enhanceTable(tableData, idx)
                tables.append(enhancedTable)

        result = {
            "success": True,
            "format": "docx",
            "fileName": path.name,
            "filePath": str(path.absolute()),
            "fileSize": path.stat().st_size,
            "content": "\n\n".join(paragraphs),
            "paragraphs": paragraphs,
            "paragraphCount": len(paragraphs),
            "tables": tables,
            "tableCount": len(tables)
        }

        logger.info(f"解析完成: {len(paragraphs)} 段落, {len(tables)} 表格")
        return result

    def _enhanceTable(self, tableData: List[List[str]], tableIndex: int) -> Dict[str, Any]:
        """增强表格信息，提取表头和关键数据"""
        if not tableData:
            return {"rawData": [], "headers": [], "rowCount": 0}

        # 尝试识别表头（第一行）
        headers = tableData[0] if tableData else []

        # 提取数值型数据（用于配置清单等）
        numericData = []
        for row in tableData[1:]:
            for cell in row:
                # 检测是否包含数字
                if any(c.isdigit() for c in cell):
                    numericData.append(row)
                    break

        return {
            "tableIndex": tableIndex,
            "rawData": tableData,
            "headers": headers,
            "rowCount": len(tableData),
            "hasNumericData": len(numericData) > 0,
            "numericRows": numericData[:10] if numericData else []  # 最多返回10行
        }

    def _parseDoc(self, path: Path) -> Dict[str, Any]:
        """解析 doc 文件 (旧格式)"""
        try:
            from docx import Document
        except ImportError:
            logger.error("缺少 python-docx 库")
            return {
                "success": False,
                "error": "请安装 python-docx 库: pip install python-docx",
                "format": "doc"
            }

        try:
            doc = Document(str(path))
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

            logger.info(f"doc 文件解析完成: {len(paragraphs)} 段落")

            return {
                "success": True,
                "format": "doc",
                "fileName": path.name,
                "filePath": str(path.absolute()),
                "fileSize": path.stat().st_size,
                "content": "\n\n".join(paragraphs),
                "paragraphs": paragraphs,
                "paragraphCount": len(paragraphs),
                "note": "doc 格式支持有限，建议转换为 docx 格式以获得更好的解析效果"
            }
        except Exception as e:
            logger.error(f"解析 doc 文件失败: {e}")
            return {
                "success": False,
                "error": f"解析 doc 文件失败: {str(e)}",
                "format": "doc",
                "suggestion": "请将 doc 文件转换为 docx 格式后重试"
            }

    def _parsePdf(self, path: Path) -> Dict[str, Any]:
        """解析 PDF 文件"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("缺少 PyPDF2 库")
            return {
                "success": False,
                "error": "请安装 PyPDF2 库: pip install PyPDF2",
                "format": "pdf"
            }

        try:
            reader = PdfReader(str(path))
            logger.debug(f"PDF 文件总页数: {len(reader.pages)}")

            pages = []
            fullText = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append({
                        "pageNumber": i + 1,
                        "content": text.strip(),
                        "charCount": len(text)
                    })
                    fullText.append(text.strip())

            logger.info(f"PDF 解析完成: {len(pages)} 页")

            return {
                "success": True,
                "format": "pdf",
                "fileName": path.name,
                "filePath": str(path.absolute()),
                "fileSize": path.stat().st_size,
                "content": "\n\n".join(fullText),
                "pages": pages,
                "pageCount": len(pages),
                "totalCharCount": sum(p["charCount"] for p in pages)
            }
        except Exception as e:
            logger.error(f"解析 PDF 文件失败: {e}")
            return {
                "success": False,
                "error": f"解析 PDF 文件失败: {str(e)}",
                "format": "pdf"
            }

    def parseToJson(self, filePath: str, outputPath: Optional[str] = None) -> str:
        """
        解析文档并输出为 JSON 格式

        Args:
            filePath: 文档文件路径
            outputPath: 可选的输出文件路径

        Returns:
            JSON 格式的解析结果
        """
        result = self.parse(filePath)
        jsonStr = json.dumps(result, ensure_ascii=False, indent=2)

        if outputPath:
            with open(outputPath, 'w', encoding='utf-8') as f:
                f.write(jsonStr)
            logger.info(f"结果已保存到: {outputPath}")

        return jsonStr


def parseDocument(filePath: str, verbose: bool = False) -> Dict[str, Any]:
    """
    便捷函数：解析文档文件

    Args:
        filePath: 文档文件路径
        verbose: 是否显示详细日志

    Returns:
        包含解析结果的字典
    """
    parser = documentParser(verbose=verbose)
    return parser.parse(filePath)


def parseDocumentToJson(filePath: str, outputPath: Optional[str] = None, verbose: bool = False) -> str:
    """
    便捷函数：解析文档文件并返回 JSON 字符串

    Args:
        filePath: 文档文件路径
        outputPath: 可选的输出文件路径
        verbose: 是否显示详细日志

    Returns:
        JSON 格式的解析结果
    """
    parser = documentParser(verbose=verbose)
    return parser.parseToJson(filePath, outputPath)


def main():
    """命令行入口"""
    # 设置 stdout 编码为 UTF-8，避免 Windows 控制台乱码
    if sys.platform == 'win32':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except Exception:
            pass

    parser = argparse.ArgumentParser(
        description='文档解析器 - 解析 doc, docx, pdf 文件',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python parser.py contract.docx
  python parser.py contract.docx -o result.json
  python parser.py contract.pdf --format json -v
        '''
    )

    parser.add_argument('file', nargs='?', help='要解析的文档文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径 (可选)')
    parser.add_argument('-f', '--format', choices=['json', 'text'], default='json',
                        help='输出格式 (默认: json)')
    parser.add_argument('-v', '--verbose', action='store_true',
                        help='显示详细日志')

    args = parser.parse_args()

    if not args.file:
        parser.print_help()
        return

    try:
        docParser = documentParser(verbose=args.verbose)
        result = docParser.parse(args.file)

        if args.format == 'json':
            output = json.dumps(result, ensure_ascii=False, indent=2)
        else:
            # 纯文本格式
            output = result.get('content', '')

        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(output)
            print(f"解析结果已保存到: {args.output}")
        else:
            print(output)

    except FileNotFoundError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"解析失败: {e}", file=sys.stderr)
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
